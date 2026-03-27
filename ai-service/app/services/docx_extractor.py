from docx import Document

def extract(file_path: str) -> str:
    """Extracts text from a DOCX file."""
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    full_text.append(" ".join(row_data))
                    
        return "\n".join(full_text)
    except Exception as e:
        return f"Error: {str(e)}"
